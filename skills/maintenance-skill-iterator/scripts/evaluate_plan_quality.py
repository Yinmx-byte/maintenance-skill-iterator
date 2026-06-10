from __future__ import annotations

import argparse
import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document


CANONICAL_HEADINGS = [
    "背景",
    "检修类型",
    "现场环境",
    "实施计划",
    "4.1 检修窗口",
    "4.2 实施人员",
    "风险评估",
    "5.1影响范围",
    "5.2危险点分析",
    "5.3安全措施",
    "5.3.1授权",
    "5.3.2备份",
    "5.3.3验证",
    "5.3.4 双人复核",
    "实施步骤",
    "6.1备份",
    "6.2 检修前验证",
    "6.3 检修操作",
    "6.4 检修后验证",
    "回滚步骤",
    "7.1 回滚操作",
    "7.2 回滚后验证",
]

GENERIC_PHRASES = [
    "根据实际情况",
    "视情况",
    "按需",
    "相关人员",
    "确保正常",
    "进行检查",
    "完成后验证",
]

ACTION_ALIASES = {
    "create": ["create", "创建", "新建", "申请"],
    "recycle": ["recycle", "回收", "释放", "删除", "空闲"],
    "resize": ["resize", "升配", "降配", "规格变更", "扩容", "缩容"],
    "restart": ["restart", "重启", "维护性重启"],
    "drill": ["drill", "演练", "单点", "切流"],
    "ipv6": ["ipv6", "IPv6"],
}


@dataclass
class DocProfile:
    path: str
    headings: list[str]
    paragraphs: list[str]
    table_count: int

    @property
    def text(self) -> str:
        return "\n".join(self.paragraphs)


def normalize(text: str) -> str:
    return re.sub(r"\s+", "", text or "").lower()


def read_docx(path: Path) -> DocProfile:
    doc = Document(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    headings: list[str] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style = p.style.name if p.style else ""
        if (
            style.startswith("Heading")
            or text[:2] in {"一、", "二、", "三、", "四、", "五、", "六、", "七、"}
            or re.match(r"^\d+(\.\d+)+\s*", text)
            or text in {"背景", "检修类型", "现场环境", "实施计划", "风险评估", "实施步骤", "回滚步骤"}
        ):
            headings.append(text)
    return DocProfile(str(path), headings, paragraphs, len(doc.tables))


def read_references(reference_dir: Path) -> list[DocProfile]:
    docs = []
    for path in sorted(reference_dir.glob("*.docx")):
        if path.name.startswith("~$"):
            continue
        try:
            docs.append(read_docx(path))
        except Exception:
            continue
    if not docs:
        raise SystemExit(f"No .docx reference files found in {reference_dir}")
    return docs


def text_from_candidate(args: argparse.Namespace) -> tuple[str, dict]:
    if args.candidate_docx:
        profile = read_docx(Path(args.candidate_docx))
        return profile.text, {
            "type": "docx",
            "path": profile.path,
            "headings": profile.headings,
            "tables": profile.table_count,
            "source_skill": args.source_skill or args.candidate_skill,
            "reference_dir": args.reference_dir,
        }
    if args.candidate_skill:
        path = Path(args.candidate_skill)
        return path.read_text(encoding="utf-8"), {"type": "skill", "path": str(path)}
    raise SystemExit("Provide --candidate-skill or --candidate-docx")


def coverage(items: Iterable[str], text: str) -> tuple[int, list[str]]:
    normalized_text = normalize(text)
    missing = [item for item in items if normalize(item) not in normalized_text]
    total = len(list(items)) if not isinstance(items, list) else len(items)
    score = round((total - len(missing)) * 100 / total) if total else 100
    return score, missing


def keyword_score(groups: dict[str, list[str]], text: str) -> tuple[int, dict[str, list[str]]]:
    missing_by_group: dict[str, list[str]] = {}
    scores = []
    for group, words in groups.items():
        score, missing = coverage(words, text)
        scores.append(score)
        if missing:
            missing_by_group[group] = missing
    return round(sum(scores) / len(scores)), missing_by_group


def extract_frontmatter_name(text: str, fallback: str = "unknown") -> str:
    match = re.search(r"^---\s*(.*?)\s*---", text, flags=re.S)
    if not match:
        return fallback
    name_match = re.search(r"^name:\s*([^\n]+)", match.group(1), flags=re.M)
    return name_match.group(1).strip().strip("'\"") if name_match else fallback


def extract_quality_contract(skill_path: str | None) -> dict:
    if not skill_path:
        return {"product": "unknown", "checks": {}, "found": False}
    path = Path(skill_path)
    if not path.exists():
        return {"product": path.stem or "unknown", "checks": {}, "found": False}
    text = path.read_text(encoding="utf-8-sig")
    product = extract_frontmatter_name(text, path.parent.name)
    heading = re.search(r"^##\s*质量评估契约\s*$", text, flags=re.M)
    if not heading:
        return {"product": product, "checks": {}, "found": False}
    rest = text[heading.end() :]
    next_heading = re.search(r"^##\s+", rest, flags=re.M)
    contract_text = rest[: next_heading.start()] if next_heading else rest
    checks: dict[str, list[str]] = {}
    current_group = "common"
    checks[current_group] = []
    for raw_line in contract_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        section_match = re.match(r"^###\s+(.+)$", line)
        if section_match:
            current_group = normalize_contract_group(section_match.group(1))
            checks.setdefault(current_group, [])
            continue
        bullet_match = re.match(r"^[-*]\s+(.+)$", line)
        if bullet_match:
            checks.setdefault(current_group, []).extend(split_contract_item(bullet_match.group(1)))
    cleaned = {
        group: dedupe([item for item in items if item])
        for group, items in checks.items()
        if any(item for item in items)
    }
    return {"product": product, "checks": cleaned, "found": bool(cleaned)}


def normalize_contract_group(title: str) -> str:
    title_norm = normalize(title)
    for group, aliases in ACTION_ALIASES.items():
        if any(normalize(alias) in title_norm for alias in aliases):
            return group
    if "通用" in title or "common" in title_norm:
        return "common"
    return re.sub(r"[^a-zA-Z0-9_\-\u4e00-\u9fff]+", "_", title.strip()).strip("_").lower()


def split_contract_item(item: str) -> list[str]:
    text = re.sub(r"^(必须|应|需要|检查|包含|出现|写明|覆盖)[：:]\s*", "", item.strip())
    text = re.sub(r"^(必须|应|需要)(包含|出现|写明|覆盖)[：:]\s*", "", text)
    if "：" in text or ":" in text:
        prefix, rest = re.split(r"[：:]", text, maxsplit=1)
        if any(word in prefix for word in ["必须", "包含", "出现", "字段", "检查项"]):
            text = rest
    parts = re.split(r"[、,，；;]", text)
    normalized = [part.strip(" `。；;，,") for part in parts]
    return [part for part in normalized if part]


def dedupe(items: list[str]) -> list[str]:
    seen = set()
    result = []
    for item in items:
        key = normalize(item)
        if key and key not in seen:
            seen.add(key)
            result.append(item)
    return result


def select_contract_groups(contract: dict, candidate_meta: dict, candidate_text: str) -> dict[str, list[str]]:
    checks = contract.get("checks") or {}
    if candidate_meta.get("type") == "skill":
        return checks
    selected = {}
    if checks.get("common"):
        selected["common"] = checks["common"]
    action_hint = normalize(
        str(candidate_meta.get("reference_dir", ""))
        + "\n"
        + "\n".join(candidate_meta.get("headings", [])[:8])
    )
    fallback_hint = normalize(candidate_text[:1000])
    for group in ["create", "recycle", "resize", "restart", "drill"]:
        aliases = ACTION_ALIASES[group]
        if group in checks and any(normalize(alias) in action_hint for alias in aliases):
            selected[group] = checks[group]
            break
    if "ipv6" in checks and any(normalize(alias) in fallback_hint for alias in ACTION_ALIASES["ipv6"]):
        selected["ipv6"] = checks["ipv6"]
    if len(selected) == (1 if "common" in selected else 0):
        for group in ["create", "recycle", "resize", "restart", "drill"]:
            aliases = ACTION_ALIASES[group]
            if group in checks and any(normalize(alias) in fallback_hint for alias in aliases):
                selected[group] = checks[group]
                break
    return selected or checks


def build_findings(
    heading_missing: list[str],
    contract_missing: dict[str, list[str]],
    generic_hits: list[str],
    candidate_meta: dict,
    product: str,
    duplicate_numbering: bool,
    contract_found: bool,
) -> list[dict]:
    findings = []
    is_generated_docx = candidate_meta.get("type") == "docx"
    target_name = "生成结果 DOCX" if is_generated_docx else "候选 Skill"
    if heading_missing:
        findings.append(
            {
                "severity": "high",
                "category": "structure",
                "message": f"{target_name}未覆盖参考方案的稳定章节：" + "、".join(heading_missing[:10]),
                "suggested_skill_change": "在源 Skill 中明确要求生成结果保留参考方案的固定章节和编号，尤其是风险评估、实施步骤、回滚步骤下的二级/三级子项。",
            }
        )
    if not contract_found:
        findings.append(
            {
                "severity": "medium",
                "category": "contract",
                "message": "源 Skill 未提供 `## 质量评估契约`，只能执行通用章节、风险、步骤和格式检查。",
                "suggested_skill_change": "在源 Skill 中补充 `## 质量评估契约`，用 `### common` 和动作小节列出生成结果必须覆盖的检查项。",
            }
        )
    for group, missing in contract_missing.items():
        severity = "high" if group in {"common", "create", "recycle", "drill"} else "medium"
        findings.append(
            {
                "severity": severity,
                "category": "contract",
                "message": f"{target_name}未满足 {product} / {group} 质量评估契约，缺少：" + "、".join(missing),
                "suggested_skill_change": f"在源 Skill 的 {group} 生成规则中强化这些检查项，要求模型必须写入生成 DOCX。",
            }
        )
    if generic_hits:
        findings.append(
            {
                "severity": "medium",
                "category": "operation",
                "message": f"{target_name}包含空泛措辞：" + "、".join(generic_hits),
                "suggested_skill_change": "在源 Skill 中禁止这些空泛措辞，并要求实施步骤写到控制台路径、对象名称、配置项、验证动作和预期结果。",
            }
        )
    if duplicate_numbering:
        findings.append(
            {
                "severity": "medium",
                "category": "format",
                "message": f"{target_name}存在重复编号，例如“1、1、”。",
                "suggested_skill_change": "在通用 composer Skill 中明确要求 numbered_list 的 items 不得自带“1、”“2、”等序号，或在渲染器中对编号前缀做清洗。",
            }
        )
    if candidate_meta.get("type") == "docx" and candidate_meta.get("tables", 0) < 2:
        findings.append(
            {
                "severity": "medium",
                "category": "format",
                "message": f"候选 DOCX 表格数量少于参考 {product.upper()} 方案常见数量。",
                "suggested_skill_change": "在源 Skill 或通用 composer Skill 中要求生成结果输出资源信息、实施人员/窗口、风险或操作清单等表格化内容。",
            }
        )
    return findings


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--reference-dir", required=True)
    parser.add_argument("--candidate-skill")
    parser.add_argument("--candidate-docx")
    parser.add_argument("--source-skill")
    parser.add_argument("--output")
    args = parser.parse_args()

    references = read_references(Path(args.reference_dir))
    candidate_text, candidate_meta = text_from_candidate(args)

    reference_summary = {
        "count": len(references),
        "files": [
            {
                "path": doc.path,
                "paragraphs": len(doc.paragraphs),
                "tables": doc.table_count,
                "headings": doc.headings,
            }
            for doc in references
        ],
    }

    structure_score, heading_missing = coverage(CANONICAL_HEADINGS, candidate_text)
    contract = extract_quality_contract(candidate_meta.get("source_skill") or candidate_meta.get("path"))
    product = contract["product"]
    contract_groups = select_contract_groups(contract, candidate_meta, candidate_text)
    contract_score, contract_missing = keyword_score(contract_groups, candidate_text) if contract_groups else (100, {})
    operation_score, operation_missing = coverage(
        ["备份", "检修前验证", "检修操作", "检修后验证"], candidate_text
    )
    rollback_score, rollback_missing = coverage(["回滚操作", "回滚后验证"], candidate_text)
    generic_hits = [phrase for phrase in GENERIC_PHRASES if phrase in candidate_text]
    duplicate_numbering = bool(re.search(r"\b\d+、\s*\d+、", candidate_text))

    risk_score, risk_missing = coverage(
        ["影响范围", "危险点分析", "安全措施", "授权", "备份", "验证", "双人复核"],
        candidate_text,
    )
    if candidate_meta["type"] == "skill":
        if "maintenance-plan-composer" in candidate_meta.get("path", ""):
            format_score, _format_missing = coverage(
                ["JSON", "document.sections", "blocks", "表格", "首页"],
                candidate_text,
            )
        else:
            format_score = 100
    else:
        format_score = min(100, 40 + candidate_meta.get("tables", 0) * 20)
        if duplicate_numbering:
            format_score = min(format_score, 85)

    findings = build_findings(
        heading_missing + operation_missing + rollback_missing + risk_missing,
        contract_missing,
        generic_hits,
        candidate_meta,
        product,
        duplicate_numbering,
        contract["found"],
    )
    dimension_scores = {
        "structure": structure_score,
        "risk": risk_score,
        "operation": operation_score,
        "rollback": rollback_score,
        "contract": contract_score,
        "format": format_score,
    }
    score = round(sum(dimension_scores.values()) / len(dimension_scores))

    evaluation_mode = "generated_docx" if candidate_meta["type"] == "docx" else "skill_static_preflight"
    if findings:
        if evaluation_mode == "generated_docx":
            patch_summary = "生成结果存在缺陷；应将 findings 中的缺失章节、动作特异性风险、可执行实施步骤和回滚闭环要求回写到 source_skill。"
        else:
            patch_summary = "静态预检发现源 Skill 规则缺失；修复后仍需生成 DOCX 并进行 generated_docx 评估。"
    else:
        if evaluation_mode == "generated_docx":
            patch_summary = "生成结果已覆盖当前规则检查项，暂不建议自动修改源 Skill；可抽样人工复核措辞和格式。"
        else:
            patch_summary = "源 Skill 静态规则覆盖良好；下一步必须评估实际生成 DOCX 的质量。"

    result = {
        "evaluation_mode": evaluation_mode,
        "product": product,
        "contract_groups": list(contract_groups.keys()),
        "score": score,
        "dimension_scores": dimension_scores,
        "reference_summary": reference_summary,
        "candidate": candidate_meta,
        "findings": findings,
        "recommended_patch_summary": patch_summary,
    }

    output = json.dumps(result, ensure_ascii=False, indent=2)
    if args.output:
        Path(args.output).write_text(output, encoding="utf-8")
    else:
        print(output)


if __name__ == "__main__":
    main()
